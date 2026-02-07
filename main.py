from fastapi import FastAPI, UploadFile, File, Form, Request, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
import re, json, os, uuid
from groq import Groq

# ---------- APP SETUP ----------

app = FastAPI()
client = Groq()  # uses GROQ_API_KEY

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")
templates = Jinja2Templates(directory=os.path.join(BASE_DIR, "templates"))

UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------- HELPERS ----------

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    return "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

def split_into_clauses(text: str):
    pattern = r"(?:\n|\r)(\d+\.\s+[A-Z][^\n]+)"
    parts = re.split(pattern, text)

    clauses = []
    title = "Introduction"
    body = ""

    for part in parts:
        if re.match(r"\d+\.\s+", part):
            if body.strip():
                clauses.append({"title": title, "text": body.strip()})
            title = part.strip()
            body = ""
        else:
            body += "\n" + part

    if body.strip():
        clauses.append({"title": title, "text": body.strip()})

    return clauses

def apply_and_highlight(text: str, changes: list) -> str:
    updated = text
    applied = False

    for c in changes:
        original = c.get("original_excerpt", "").strip()
        revised = c.get("revised_text", "").strip()

        if not original or not revised:
            continue

        if original in updated:
            updated = updated.replace(
                original,
                f'<mark class="ai-change">{revised}</mark>',
            )
            applied = True

    return updated

def save_docx(html_text: str) -> str:
    clean = html_text.replace("<mark>", "").replace("</mark>", "")
    doc = Document()

    for block in clean.split("\n\n"):
        doc.add_paragraph(block)

    filename = f"{uuid.uuid4()}.docx"
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return filename

# ---------- ROUTES ----------

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/preview")
async def preview_contract(file: UploadFile = File(...)):
    path = os.path.join(UPLOAD_DIR, file.filename)
    with open(path, "wb") as f:
        f.write(await file.read())

    return {"contract_text": extract_text_from_docx(path)}

@app.post("/process", response_class=HTMLResponse)
async def process_changes(
    request: Request,
    contract_text: str = Form(...),
    change_instructions: str = Form(...)
):
    clauses = split_into_clauses(contract_text)

    prompt = f"""
You are a professional contract editor.

RULES:
- Identify ALL clauses impacted by the user request
- Apply changes to EVERY relevant occurrence
- If the same term appears multiple times, update ALL instances
- Do NOT invent clauses unless explicitly requested
- original_excerpt must match text EXACTLY
- Prefer minimal, precise legal edits
- You MUST search the ENTIRE contract for every place the change applies
- If the same sentence or concept appears multiple times, you MUST return a SEPARATE change entry for EACH occurrence

Return STRICT JSON ONLY:
{{
  "changes": [
    {{
      "clause_title": "...",
      "original_excerpt": "...",
      "revised_text": "...",
      "reason": "..."
    }}
  ]
}}

User request:
{change_instructions}

Clauses:
{json.dumps(clauses)}
"""

    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[
            {"role": "system", "content": "Return valid JSON only."},
            {"role": "user", "content": prompt}
        ],
        temperature=0
    )

    ai_data = json.loads(response.choices[0].message.content)
    changes = ai_data.get("changes", [])

    highlighted = apply_and_highlight(contract_text, changes)
    docx_name = save_docx(highlighted)

    return templates.TemplateResponse(
        "result.html",
        {
            "request": request,
            "contract": highlighted,
            "download_file": docx_name
        }
    )

@app.post("/download-edit", response_class=FileResponse)
async def download_edited(text: str = Form(...)):
    filename = save_docx(text)
    path = os.path.join(OUTPUT_DIR, filename)

    return FileResponse(
        path=path,
        filename="revised_contract.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
